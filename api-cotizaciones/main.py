import sys, os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))


import re
import unicodedata
from typing import Any, Dict
import psycopg2
from psycopg2.extras import execute_values, Json

from fastapi import FastAPI, UploadFile, File, HTTPException, Request, Header
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import Optional, List, Tuple
from typing import Dict
from decimal import Decimal, getcontext
from docx import Document
from datetime import datetime, date
from zoneinfo import ZoneInfo
import json
import uuid
import requests
import subprocess
from utils.parser import extraer_variables
from dotenv import load_dotenv

from datetime import datetime, timedelta



# -------------------------------------------------
# CONFIGURACIÓN GENERAL
# -------------------------------------------------
getcontext().prec = 28

app = FastAPI(
    title="API Unificada de Cotización y Documentos",
    description="Servicio que calcula cotizaciones, administra plantillas y genera documentos Word/PDF."
)

TEMPLATES_DIR = "templates"
OUTPUT_DIR = "outputs"
DB_PATH = "db.json"
TIMEZONE = ZoneInfo("America/Mexico_City")

os.makedirs(TEMPLATES_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

load_dotenv()
# API KEY para actualizar variables
API_ADMIN_KEY = os.getenv("API_ADMIN_KEY")
if not API_ADMIN_KEY:
    raise RuntimeError(
        "❌ ERROR: Debes definir la variable de entorno API_ADMIN_KEY (en .env o en el entorno del servidor)."
    )

# -------------------------------------------------
# SYNC CARTERA (Postgres en Render)
# -------------------------------------------------
CARTERA_DATABASE_URL = os.getenv("CARTERA_DATABASE_URL")
CARTERA_SYNC_API_KEY = os.getenv("CARTERA_SYNC_API_KEY")

# -------------------------------------------------
# VARIABLES / PARÁMETROS DE NEGOCIO (por defecto)
# -------------------------------------------------
DEFAULT_VARIABLES = {
    "planes": {
        "CLASICO": {
            "div_plan": 40,
            "seguro_anual": -1,
            "seguro_contado": True,
            "segurotag": "Seguro de Daños",
            "GPStag": "GPS",
            "plantag": "Gastos Administrativos",
            "incluye_asistencia":False,
            "incluye_gestoria": True,
            "incluye_sustituto": False,
            "incluye_GPS": True,
        },
        "PLUS": {
            "div_plan": 44,
            "seguro_anual": -1,
            "seguro_contado": False,
            "segurotag": "Seguro con Deducibles 5 y 10%",
            "GPStag": "GPS con Plataforma de Rastreo",
            "plantag": "Membresia Plan Plus",
            "incluye_asistencia":False,
            "incluye_gestoria": True,
            "incluye_sustituto": False,
            "incluye_GPS": True,
        },
        "PREMIUM": {
            "div_plan": 48,
            "seguro_anual": -1,
            "seguro_contado": False,
            "segurotag": "Seguro Premium Deducibles 3 y 5%",
            "GPStag": "GPS con Plataforma de Rastreo",
            "plantag": "Membresia Plan Premium",
            "incluye_asistencia":True,
            "incluye_gestoria": True,
            "incluye_sustituto": True,
            "incluye_GPS": True,
        },
    },
    "defaults": {
        "tasa_anual_default": 27.0,
        "enganche_default": 10.0,
        "rentas_deposito_default": 0.0,
        "comision_default": 3.0,
        "valor_default": 0.0,

        "gestoria_default": 2320,
        "localizador_anual_default": 1976.64,
        "localizador_inicial_default": 4491.52,

        "residuales_default": [
            {"plazo": 24, "residual": 40},
            {"plazo": 36, "residual": 30},
            {"plazo": 48, "residual": 25},
            {"plazo": 60, "residual": 20},
        ],
        "seguro_por_monto": [
            {"max_valor_con_iva": 499999, "porcentaje": 0.04},
            {"max_valor_con_iva": 749999, "porcentaje": 0.035},
            {"max_valor_con_iva": 999999, "porcentaje": 0.03},
            {"max_valor_con_iva": 1499999, "porcentaje": 0.0275},
            {"max_valor_con_iva": 4999999, "porcentaje": 0.025},
            {"max_valor_con_iva": 9999999999, "porcentaje": 0.025},
        ],
    },
}

VARIABLES = DEFAULT_VARIABLES.copy()

def _plan_key(plan: Optional[str]) -> str:
    p = (plan or "PREMIUM").strip().upper()
    return p if p else "PREMIUM"

def get_defaults() -> dict:
    d = VARIABLES.get("defaults")
    return d if isinstance(d, dict) else DEFAULT_VARIABLES["defaults"]

def get_planes() -> dict:
    p = VARIABLES.get("planes")
    return p if isinstance(p, dict) else DEFAULT_VARIABLES["planes"]

def get_plan_cfg(plan: Optional[str]) -> dict:
    planes = get_planes()
    key = _plan_key(plan)
    return planes.get(key) or planes.get("PREMIUM") or DEFAULT_VARIABLES["planes"]["PREMIUM"]
# -------------------------------------------------
# PLANTILLA PRINCIPAL (carga automática desde GitHub)
# -------------------------------------------------
GITHUB_RAW_URL = (
    "https://github.com/FirstLeaseAgent/api-cotizaciones/raw/refs/heads/main/"
    "api-cotizaciones/templates/Plantilla_Cotizacion.docx"
)
TEMPLATE_NAME = "Plantilla_Cotizacion.docx"


def ensure_db_and_variables():
    """
    Asegura que db.json exista y que variables tenga la forma:
      variables: { "planes": {...}, "defaults": {...} }
    """
    global VARIABLES

    if not os.path.exists(DB_PATH) or os.stat(DB_PATH).st_size == 0:
        with open(DB_PATH, "w") as f:
            json.dump({"plantillas": [], "variables": DEFAULT_VARIABLES}, f, indent=4)
        VARIABLES = DEFAULT_VARIABLES.copy()
        return

    with open(DB_PATH, "r+") as f:
        data = json.load(f)
        changed = False

        if "plantillas" not in data:
            data["plantillas"] = []
            changed = True

        if "variables" not in data or not isinstance(data["variables"], dict):
            data["variables"] = DEFAULT_VARIABLES
            changed = True

        vars_db = data["variables"]

        # --- Migración: si venía estructura vieja/plana, crear defaults/planes ---
        if "defaults" not in vars_db or "planes" not in vars_db:
            # Si ya había algo plano, lo intentamos meter a defaults (solo las llaves que aplican)
            legacy = vars_db.copy()

            vars_db = {
                "planes": legacy.get("planes") if isinstance(legacy.get("planes"), dict) else DEFAULT_VARIABLES["planes"],
                "defaults": legacy.get("defaults") if isinstance(legacy.get("defaults"), dict) else DEFAULT_VARIABLES["defaults"],
            }

            # Si el legacy traía llaves planas típicas, las mapeamos a defaults
            for k in ("tasa_anual_default","enganche_default","rentas_deposito_default","comision_default","valor_default"):
                if k in legacy and legacy[k] is not None:
                    vars_db["defaults"][k] = legacy[k]

            if "residuales_default" in legacy and legacy["residuales_default"] is not None:
                vars_db["defaults"]["residuales_default"] = legacy["residuales_default"]

            if "seguro_por_monto" in legacy and legacy["seguro_por_monto"] is not None:
                vars_db["defaults"]["seguro_por_monto"] = legacy["seguro_por_monto"]

            changed = True

        # Asegurar que existan defaults/planes y llaves base
        if not isinstance(vars_db.get("planes"), dict):
            vars_db["planes"] = DEFAULT_VARIABLES["planes"]
            changed = True
        if not isinstance(vars_db.get("defaults"), dict):
            vars_db["defaults"] = DEFAULT_VARIABLES["defaults"]
            changed = True

        # Completar faltantes en defaults
        for k, v in DEFAULT_VARIABLES["defaults"].items():
            if k not in vars_db["defaults"]:
                vars_db["defaults"][k] = v
                changed = True

        # Completar faltantes en planes (sin pisar los existentes)
        for plan_name, cfg in DEFAULT_VARIABLES["planes"].items():
            if plan_name not in vars_db["planes"] or not isinstance(vars_db["planes"][plan_name], dict):
                vars_db["planes"][plan_name] = cfg
                changed = True
            else:
                for k, v in cfg.items():
                    if k not in vars_db["planes"][plan_name]:
                        vars_db["planes"][plan_name][k] = v
                        changed = True

        data["variables"] = vars_db
        if changed:
            f.seek(0)
            f.truncate()
            json.dump(data, f, indent=4)

        VARIABLES = data["variables"]


def ensure_template_available():
    ensure_db_and_variables()

    template_path = os.path.join(TEMPLATES_DIR, TEMPLATE_NAME)

    # Descarga si no existe la plantilla
    if not os.path.exists(template_path):
        print("🔄 Descargando plantilla desde GitHub…")
        resp = requests.get(GITHUB_RAW_URL)
        resp.raise_for_status()
        with open(template_path, "wb") as f:
            f.write(resp.content)
        print("✅ Plantilla descargada.")

    # ✅ Detectar variables del DOCX
    detected_vars = extraer_variables(template_path)


    # Registrar/actualizar plantilla en DB
    with open(DB_PATH, "r+") as f:
        data = json.load(f)

        # Si no hay plantillas, crear
        if not data.get("plantillas"):
            plantilla_id = str(uuid.uuid4())
            data["plantillas"] = [{
                "id": plantilla_id,
                "nombre": TEMPLATE_NAME,
                "variables": detected_vars
            }]
            print("✅ Plantilla registrada en db.json (con variables detectadas)")
        else:
            # Si ya hay, actualiza la primera (o busca por nombre)
            updated = False
            for p in data["plantillas"]:
                if p.get("nombre") == TEMPLATE_NAME:
                    p["variables"] = detected_vars
                    updated = True
                    break
            if not updated:
                data["plantillas"].append({
                    "id": str(uuid.uuid4()),
                    "nombre": TEMPLATE_NAME,
                    "variables": detected_vars
                })
            print("✅ Variables de plantilla actualizadas en db.json")

        f.seek(0)
        f.truncate()
        json.dump(data, f, indent=4, ensure_ascii=False)


ensure_template_available()


# -------------------------------------------------
# MODELOS Pydantic
# -------------------------------------------------
class ResidualItem(BaseModel):
    plazo: int
    residual: float


class CotizacionRequest(BaseModel):
    nombre: str
    nombre_activo: str
    valor: float

    # NUEVO: Plan (opcional, default PREMIUM)
    plan: Optional[str] = None

    # Defaults globales (si None -> defaults)
    enganche: Optional[float] = None
    tasa_anual: Optional[float] = None
    comision: Optional[float] = None
    rentas_deposito: Optional[float] = None

    # NUEVO: div_plan por request (si None -> del plan)
    div_plan: Optional[float] = None

    # Gestoría por request (si None -> del plan)
    gestoria: Optional[float] = None

    # Seguro:
    #   None o -1 -> calcular por tabla
    #   0 -> sin seguro
    #   >0 -> monto anual SIN IVA (como tu lógica actual)
    seguro_anual: Optional[float] = None

    # Si no viene, se toma del plan
    seguro_contado: Optional[bool] = None

    accesorios: Optional[float] = 0.0              # con IVA
    localizador_inicial: Optional[float] = None     # con IVA
    localizador_anual: Optional[float] = None       # con IVA

    residuales: Optional[List[ResidualItem]] = None

    incluye_gestoria: Optional[bool] = None
    incluye_sustituto: Optional[bool] = None
    incluye_GPS: Optional[bool] = None
    incluye_asistencia: Optional[bool] = None
# ------------------------------------------
# ACTUALIZAR EJEMPLO DE SWAGGER DINÁMICAMENTE
# ------------------------------------------
def _vdef(key: str, fallback=None):
    d = VARIABLES.get("defaults")
    if isinstance(d, dict) and key in d:
        return d[key]
    # fallback legacy
    if key in VARIABLES:
        return VARIABLES[key]
    if fallback is not None:
        return fallback
    return DEFAULT_VARIABLES.get(key)

def refresh_cotizar_example():
    D = get_defaults()
    P = get_plan_cfg("PREMIUM")

    ejemplo = {
        "plan": "PREMIUM",
        "nombre": "Cliente Ejemplo",
        "nombre_activo": "Camioneta Tiguan 2025",
        "valor": 0,
        "enganche": D["enganche_default"],
        "tasa_anual": D["tasa_anual_default"],
        "comision": D["comision_default"],
        "rentas_deposito": D["rentas_deposito_default"],
        "seguro_anual": -1,
        "seguro_contado": P.get("seguro_contado", False),
        "accesorios": 0,
        "localizador_inicial": P.get("localizador_inicial", 0),
        "localizador_anual": P.get("localizador_anual", 0),
        "residuales": D["residuales_default"],
        "gestoria": P.get("gestoria", 0),
        "div_plan": P.get("div_plan", 48),
        "incluye_sustituto": True,
        "incluye_gestoria": True,
        "incluye_GPS": True,
        "incluye_asistencia": P.get("incluye_asistencia", False),

    }

    CotizacionRequest.model_config["json_schema_extra"] = {"example": ejemplo}


# Ejecutar tras definir la clase
refresh_cotizar_example()

class SeguroRango(BaseModel):
    max_valor_con_iva: float
    porcentaje: float


class ResidualConfig(BaseModel):
    plazo: int
    residual: float


class PlanUpdate(BaseModel):
    div_plan: Optional[float] = None
    seguro_anual: Optional[float] = None
    seguro_contado: Optional[bool] = None
    segurotag: Optional[str] = None
    GPStag: Optional[str] = None
    plantag: Optional[str] = None
    incluye_gestoria: Optional[bool] = None
    incluye_sustituto: Optional[bool] = None
    incluye_GPS: Optional[bool] = None
    incluye_asistencia: Optional[bool] = None


class DefaultsUpdate(BaseModel):
    tasa_anual_default: Optional[float] = None
    enganche_default: Optional[float] = None
    rentas_deposito_default: Optional[float] = None
    comision_default: Optional[float] = None
    valor_default: Optional[float] = None
    residuales_default: Optional[List[ResidualConfig]] = None
    seguro_por_monto: Optional[List[SeguroRango]] = None
    gestoria_default: Optional[float] = None
    localizador_anual_default: Optional[float] = None
    localizador_inicial_default: Optional[float] = None

class VariablesUpdate(BaseModel):
    planes: Optional[Dict[str, PlanUpdate]] = None
    defaults: Optional[DefaultsUpdate] = None


# -------------------------------------------------
# SEGURO ANUAL
# -------------------------------------------------
def calcular_seguro_anual(valor_con_iva: float, entrada: Optional[float]) -> Decimal:
    """
    Calcula el seguro anual SIN IVA.

    - Si entrada es None o -1 → usa tabla VARIABLES["seguro_por_monto"] (basada en VALOR CON IVA).
    - Si entrada es 0 → seguro gratuito.
    - Si entrada > 0 → se usa el valor proporcionado (ya SIN IVA).
    """
    # Caso: usar tabla
    if entrada is None or entrada == -1:
        v_con_iva = Decimal(str(valor_con_iva))
        valor_sin_iva = v_con_iva / Decimal("1.16")

        rangos = get_defaults().get("seguro_por_monto", DEFAULT_VARIABLES["defaults"]["seguro_por_monto"])
        pct = Decimal("0.025")
        for rango in rangos:
            max_v = Decimal(str(rango["max_valor_con_iva"]))
            if v_con_iva <= max_v:
                pct = Decimal(str(rango["porcentaje"]))
                break

        return (valor_sin_iva * pct).quantize(Decimal("0.01"))

    # Caso: seguro gratuito
    if entrada == 0:
        return Decimal("0")

    # Caso: monto proporcionado (viene CON IVA)
    return (Decimal(str(entrada)) / Decimal("1.16")).quantize(Decimal("0.01"))


# -------------------------------------------------
# CÁLCULO FINANCIERO COMPLETO
# -------------------------------------------------
def calcular_pago_mensual(
    valor,
    enganche,
    tasa_anual,
    plazo_meses,
    valor_residual,
    comision,
    rentas_deposito,
    seguro_anual,
    seguro_contado_flag,
    div_plan_pct,
    gestoria,
    accesorios_con_iva=0.0,
    localizador_inicial_con_iva=0.0,
    localizador_anual_con_iva=0.0,
):
    """
    Incluye:
    - renta principal (con residual)
    - seguro (contado o financiado)
    - gestoría
    - accesorios (sin residual, precio con IVA)
    - localizador inicial (sin residual, financiado a n)
    - localizador anual (sin residual, financiado a 12 meses)
    """

    valor_sin_iva = Decimal(valor) / Decimal("1.16")
    acc_sin_iva = Decimal(str(accesorios_con_iva or 0)) / Decimal("1.16")
    total_sin_iva = valor_sin_iva + acc_sin_iva
    enganche_pct = Decimal(str(enganche)) / Decimal("100")
    r = Decimal(tasa_anual) / Decimal(1200)
    n = Decimal(plazo_meses)

    # Valor presente del activo (sin IVA)
    pv = valor_sin_iva * (1 - enganche_pct)
    fv = valor_sin_iva * (Decimal(valor_residual) / 100)

    # ------------------------- PAGO DE RENTA PRINCIPAL -------------------------
    if r == 0:
        pago = -(pv - fv) / n
    else:
        pago = ((pv - fv * ((1 + r) ** (-n))) * r) / (1 - (1 + r) ** (-n))

    # ------------------------- SEGURO -------------------------
    pv_seguro = seguro_anual  # ya viene sin IVA

    if seguro_contado_flag and pv_seguro > 0:
        pago_seguro = Decimal("0")
        monto_seguro_contado = pv_seguro
    else:
        n_seg = Decimal(12)
        if r == 0:
            pago_seguro = -pv_seguro / n_seg
        else:
            pago_seguro = (pv_seguro * r) / (1 - (1 + r) ** (-n_seg))
        monto_seguro_contado = Decimal("0")

    # ------------------------- GESTORÍA -------------------------
    pv_gest = Decimal(str(gestoria)) / Decimal("1.16")  # lo tratamos como valor con IVA
    if r == 0:
        pago_gestoria = -pv_gest / n
    else:
        pago_gestoria = (pv_gest * r) / (1 - (1 + r) ** (-n))

    # ------------------------- ACCESORIOS -------------------------
    # accesorios_con_iva se financia sin residual, valor presente SIN IVA
    if accesorios_con_iva and accesorios_con_iva > 0:
        pv_acc_sin_iva = (Decimal(str(accesorios_con_iva)) / Decimal("1.16")) * (1 - enganche_pct)
        if r == 0:
            pago_accesorios = -pv_acc_sin_iva / n
        else:
            pago_accesorios = (pv_acc_sin_iva * r) / (1 - (1 + r) ** (-n))
    else:
        pago_accesorios = Decimal("0")

    # ------------------------- LOCALIZADOR -------------------------
    # Inicial: n meses, sin residual
    if localizador_inicial_con_iva and localizador_inicial_con_iva > 0:
        pv_loc_ini = Decimal(str(localizador_inicial_con_iva)) / Decimal("1.16")
        if r == 0:
            pago_loc_ini = -pv_loc_ini / n
        else:
            pago_loc_ini = (pv_loc_ini * r) / (1 - (1 + r) ** (-n))
    else:
        pago_loc_ini = Decimal("0")

    # Anual: 12 meses, sin residual
    if localizador_anual_con_iva and localizador_anual_con_iva > 0:
        pv_loc_anual = Decimal(str(localizador_anual_con_iva)) / Decimal("1.16")
        n_loc = Decimal(12)
        if r == 0:
            pago_loc_anual = -pv_loc_anual / n_loc
        else:
            pago_loc_anual = (pv_loc_anual * r) / (1 - (1 + r) ** (-n_loc))
    else:
        pago_loc_anual = Decimal("0")

    pago_localizador_total = pago_loc_ini + pago_loc_anual

    # ------------------------- PAGO TOTAL MENSUAL -------------------------
    subtotal_mensual = (
        pago
        + pago_seguro
        + pago_gestoria
        + pago_accesorios
        + pago_localizador_total
    )

    # El depósito se calcula sobre la renta total (como en lógica original)
    total_mensual_con_iva = subtotal_mensual * Decimal("1.16")
    monto_deposito = Decimal(rentas_deposito) * total_mensual_con_iva
    primera_mensualidad = subtotal_mensual

    monto_enganche = total_sin_iva * enganche_pct
    pv_total = total_sin_iva * (1 - enganche_pct)   # activo+accesorios neto de enganche
    monto_comision = pv_total * (Decimal(str(comision)) / Decimal("100"))

    subtotal_inicial = (
        monto_enganche
        + monto_comision
        + monto_deposito
        + primera_mensualidad
        + monto_seguro_contado
    )

    iva_inicial = (
        monto_enganche
        + monto_comision
        + primera_mensualidad
        + monto_seguro_contado
    ) * Decimal("0.16")

    total_inicial = subtotal_inicial + iva_inicial

    # ------------------------- PLAN VS RENTA -------------------------
    div = Decimal(div_plan_pct) / Decimal(100)

    renta_plan = subtotal_mensual * div
    renta_mensual = subtotal_mensual * (1 - div)

    iva_mensual = subtotal_mensual * Decimal("0.16")
    total_mensual = subtotal_mensual * Decimal("1.16")

    # ------------------------- RESIDUAL -------------------------
    monto_residual = valor_sin_iva * (Decimal(valor_residual) / 100)
    iva_residual = monto_residual * Decimal("0.16")
    total_residual = monto_residual * Decimal("1.16")

    total_final = total_residual - monto_deposito

    return {
        "Enganche": float(round(monto_enganche, 2)),
        "Comision": float(round(monto_comision, 2)),
        "Seguro_Contado": float(round(monto_seguro_contado, 2)),
        "Renta_en_Deposito": float(round(monto_deposito, 2)),
        "Primera_Mensualidad": float(round(primera_mensualidad, 2)),
        "Subtotal_Pago_Inicial": float(round(subtotal_inicial, 2)),
        "IVA_Pago_Inicial": float(round(iva_inicial, 2)),
        "Total_Inicial": float(round(total_inicial, 2)),

        "Renta_Plan": float(round(renta_plan, 2)),
        "Renta_Mensual": float(round(renta_mensual, 2)),
        "Subtotal_Mensual": float(round(subtotal_mensual, 2)),
        "IVA_Mensual": float(round(iva_mensual, 2)),
        "Total_Mensual": float(round(total_mensual, 2)),

        "Residual": float(round(monto_residual, 2)),
        "IVA_Residual": float(round(iva_residual, 2)),
        "Total_Residual": float(round(total_residual, 2)),
        "Reembolso_Deposito": float(round(-monto_deposito, 2)),
        "Total_Final": float(round(total_final, 2)),
    }


# -------------------------------------------------
# FORMATO MILES SIN DECIMALES
# -------------------------------------------------
def formato_miles(v):
    try:
        num = int(round(float(v)))
        return f"{num:,}"
    except Exception:
        return v


# -------------------------------------------------
# CONVERTIR WORD → PDF
# -------------------------------------------------
def convertir_pdf(path_word, output_dir):
    comando = [
        "soffice",
        "--headless",
        "--convert-to", "pdf",
        "--outdir", output_dir,
        path_word
    ]

    try:
        subprocess.run(comando, check=True)
    except subprocess.CalledProcessError as e:
        raise HTTPException(500, f"Error al convertir a PDF: {e}")

    pdf_name = os.path.splitext(os.path.basename(path_word))[0] + ".pdf"
    pdf_path = os.path.join(output_dir, pdf_name)

    if not os.path.exists(pdf_path):
        raise HTTPException(500, "No se generó el archivo PDF")

    return pdf_name, pdf_path


# -------------------------------------------------
# ENDPOINT PRINCIPAL /cotizar
# -------------------------------------------------
@app.post("/cotizar")
def cotizar(data: CotizacionRequest, request: Request):
    D = get_defaults()
    plan_cfg = get_plan_cfg(data.plan)
    plan = _plan_key(data.plan)

    # Defaults globales
    enganche = data.enganche if data.enganche is not None else float(D["enganche_default"])
    tasa_anual = data.tasa_anual if data.tasa_anual is not None else float(D["tasa_anual_default"])
    comision = data.comision if data.comision is not None else float(D["comision_default"])
    rentas_deposito = data.rentas_deposito if data.rentas_deposito is not None else float(D["rentas_deposito_default"])

    accesorios = float(data.accesorios or 0.0)

    # includes: request manda; si no viene, plan manda
    incl_gest = (
        data.incluye_gestoria
        if data.incluye_gestoria is not None
        else bool(plan_cfg.get("incluye_gestoria", True))
    )

    incl_gps = (
        data.incluye_GPS
        if data.incluye_GPS is not None
        else bool(plan_cfg.get("incluye_GPS", True))
    )

    incl_sust = (
        data.incluye_sustituto
        if data.incluye_sustituto is not None
        else bool(plan_cfg.get("incluye_sustituto", False))
    )

    incluye_asistencia_effective = (
        data.incluye_asistencia
        if data.incluye_asistencia is not None
        else bool(plan_cfg.get("incluye_asistencia", False))
    )
    asistenciabool = "SI" if incluye_asistencia_effective else "NO"

    # =========================
    # GESTORÍA
    # =========================
    if not incl_gest:
        gestoria_effective = 0.0
    else:
        if data.gestoria is not None:
            gestoria_effective = float(data.gestoria)
        else:
            gestoria_effective = float(D.get("gestoria_default", 0) or 0)

    # =========================
    # GPS / LOCALIZADORES
    # =========================
    if not incl_gps:
        loc_ini = 0.0
        loc_anual = 0.0
    else:
        loc_ini = (
            float(data.localizador_inicial)
            if data.localizador_inicial is not None
            else float(D.get("localizador_inicial_default", 0) or 0)
        )

        loc_anual = (
            float(data.localizador_anual)
            if data.localizador_anual is not None
            else float(D.get("localizador_anual_default", 0) or 0)
        )

    # =========================
    # ETIQUETAS
    # =========================
    autobool = "SI" if incl_sust else "NO"
    gpsbool = "SI" if incl_gps else "NO"
    gestbool = "SI" if incl_gest else "NO"

    # div_plan: request > plan
    div_plan_effective = data.div_plan if data.div_plan is not None else float(plan_cfg.get("div_plan", 48))

    # Seguro: request > plan
    seguro_in = data.seguro_anual
    if seguro_in is None:
        seguro_in = plan_cfg.get("seguro_anual", -1)

    seguro_contado_flag = data.seguro_contado if data.seguro_contado is not None else bool(plan_cfg.get("seguro_contado", False))

    # ----------------------------
    # Labels/flags para doc
    # ----------------------------
    segbool = "SI" if (seguro_in != 0) else "NO"
    

    # Tags por plan
    plantag = plan_cfg.get("plantag", "")
    segurotag = plan_cfg.get("segurotag", "")
    GPStag = plan_cfg.get("GPStag", "")
    

    nombre_upper = data.nombre.upper()
    activo_upper = data.nombre_activo.upper()

    folio = datetime.now(TIMEZONE).strftime("%Y%m%d%H%M%S")

    # Residuales
    default_residuales = D.get("residuales_default", DEFAULT_VARIABLES["defaults"]["residuales_default"])
    if data.residuales and len(data.residuales) > 0:
        enviados = {r.plazo: r.residual for r in data.residuales}
        escenarios = []
        for item in default_residuales:
            plazo = item["plazo"]
            residual = enviados.get(plazo, item["residual"])
            escenarios.append({"plazo": plazo, "residual": residual})
    else:
        escenarios = default_residuales

    # Seguro anual (sin IVA, lógica existente)
    seguro_anual = calcular_seguro_anual(data.valor, seguro_in)

    valores_para_doc = {
        # Identidad
        "nombre": nombre_upper,
        "descripcion": activo_upper,
        "fecha": datetime.now(TIMEZONE).strftime("%d/%m/%Y"),
        "folio": folio,

        # Importes base
        "precio": formato_miles(data.valor),
        "accesorios": formato_miles(accesorios),
        "ptotal": formato_miles(data.valor + accesorios),

        # Plan
        "plan": plan,
        "plantag": plantag,
        "autobool": autobool,

        # División / estructura
        "div_plan": div_plan_effective,

        # Seguro
        "segurotag": segurotag,
        "segbool": segbool,
        "asistenciabool": asistenciabool,
        "seguro_anual": formato_miles(
            float((seguro_anual * Decimal("1.16")).quantize(Decimal("0.01")))
        ),

        # Gestoría
        "gestoria": formato_miles(gestoria_effective),
        "gestbool": gestbool,

        # GPS / Localizador
        "GPStag": GPStag,
        "gpsbool": gpsbool,
        "localizador_inicial": formato_miles(loc_ini),
        "localizador_anual": formato_miles(loc_anual),
    }

    detalle = []

    # CÁLCULO DE CADA PLAZO
    for esc in escenarios:
        calc = calcular_pago_mensual(
            valor=data.valor,
            enganche=enganche,
            tasa_anual=tasa_anual,
            plazo_meses=esc["plazo"],
            valor_residual=esc["residual"],
            comision=comision,
            rentas_deposito=rentas_deposito,
            seguro_anual=seguro_anual,
            seguro_contado_flag=seguro_contado_flag,
            div_plan_pct=div_plan_effective,
            gestoria=gestoria_effective,
            accesorios_con_iva=accesorios,
            localizador_inicial_con_iva=loc_ini,
            localizador_anual_con_iva=loc_anual,
        )

        detalle.append({"Plazo": esc["plazo"], **calc})

        p = esc["plazo"]

        # Variables Word por plazo
        valores_para_doc.update({
            f"enganche{p}": formato_miles(calc["Enganche"]),
            f"comision{p}": formato_miles(calc["Comision"]),
            f"deposito{p}": formato_miles(calc["Renta_en_Deposito"]),
            f"subinicial{p}": formato_miles(calc["Subtotal_Pago_Inicial"]),
            f"IVAinicial{p}": formato_miles(calc["IVA_Pago_Inicial"]),
            f"totalinicial{p}": formato_miles(calc["Total_Inicial"]),
            f"primermes{p}": formato_miles(calc["Primera_Mensualidad"]),
            f"segurocontado{p}": formato_miles(calc["Seguro_Contado"]),

            f"mensualidad{p}": formato_miles(calc["Renta_Mensual"]),
            f"IVAmes{p}": formato_miles(calc["IVA_Mensual"]),
            f"totalmes{p}": formato_miles(calc["Total_Mensual"]),

            # Equivalencias
            f"rentaplan{p}": formato_miles(calc["Renta_Plan"]),
            f"subtotalmes{p}": formato_miles(calc["Subtotal_Mensual"]),
            f"IVAmensual{p}": formato_miles(calc["IVA_Mensual"]),
            f"totalmensual{p}": formato_miles(calc["Total_Mensual"]),

            f"residual{p}": formato_miles(calc["Residual"]),
            f"IVAresidual{p}": formato_miles(calc["IVA_Residual"]),
            f"totalresidual{p}": formato_miles(calc["Total_Residual"]),

            f"reembolso{p}": formato_miles(calc["Reembolso_Deposito"]),
            f"totalfinal{p}": formato_miles(calc["Total_Final"]),
        })

    # GENERAR DOCUMENTO
    with open(DB_PATH, "r") as f:
        data_db = json.load(f)

    plantilla = data_db["plantillas"][0]

    word_info = generar_documento_word_local(
        plantilla_id=plantilla["id"],
        valores=valores_para_doc,
        request=request
    )

    # RESPUESTA FINAL
    return {
        "Nombre": nombre_upper,
        "Activo": activo_upper,
        "Valor": round(data.valor, 2),
        "Detalle": detalle,
        "documentos": word_info,
        "parametros": {
            "accesorios": accesorios,
            "comision": comision,
            "enganche": enganche,

            "localizador_anual": loc_anual,
            "localizador_inicial": loc_ini,

            "rentas_deposito": rentas_deposito,
            "residuales": escenarios,

            "seguro_anual": float((seguro_anual * Decimal("1.16")).quantize(Decimal("0.01"))),
            "seguro_contado": seguro_contado_flag,

            "tasa_anual": tasa_anual,

            "div_plan": div_plan_effective,
            "gestoria": gestoria_effective,

            "plan": plan,

            "segurotag": segurotag,
            "GPStag": GPStag,
            "plantag": plantag,
            "autobool": autobool,
            "asistenciabool": asistenciabool,
            "segbool": segbool,
            "gestbool": gestbool,
            "gpsbool": gpsbool
        }
    }


# -------------------------------------------------
# GENERAR WORD + PDF
# -------------------------------------------------
def generar_documento_word_local(plantilla_id: str, valores: dict, request: Request):
    with open(DB_PATH, "r") as f:
        data = json.load(f)

    plantilla = next(p for p in data["plantillas"] if p["id"] == plantilla_id)
    plantilla_path = os.path.join(TEMPLATES_DIR, plantilla["nombre"])

    doc = Document(plantilla_path)

    # Reemplazo texto en párrafos
    for p in doc.paragraphs:
        for run in p.runs:
            for var, val in valores.items():
                placeholder = f"{{{{{var}}}}}"
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(val))

    # Reemplazo texto en tablas
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        for var, val in valores.items():
                            placeholder = f"{{{{{var}}}}}"
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, str(val))

    folio = valores["folio"]
    word_name = f"FirstLease-Cotizacion-{folio}.docx"
    word_path = os.path.join(OUTPUT_DIR, word_name)
    doc.save(word_path)

    # URL base
    proto = request.headers.get("x-forwarded-proto", request.url.scheme)
    host = request.headers.get("host")
    base = f"{proto}://{host}"

    url_word = f"{base}/download_word/{word_name}"

    # Intentar PDF
    pdf_name, pdf_path = convertir_pdf(word_path, OUTPUT_DIR)
    url_pdf = f"{base}/download_pdf/{pdf_name}"

    return {
        "archivo_word": word_name,
        "descargar_word": url_word,
        "folio": folio,
        "archivo_pdf": pdf_name,
        "descargar_pdf": url_pdf
    }


# -------------------------------------------------
# ENDPOINTS AUXILIARES
# -------------------------------------------------
@app.get("/download_word/{filename}")
def download_word(filename: str):
    path = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(path):
        raise HTTPException(404, "Archivo no encontrado")
    return FileResponse(path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.get("/download_pdf/{filename}")
def download_pdf(filename: str):
    path = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(path):
        raise HTTPException(404, "Archivo no encontrado")
    return FileResponse(path, media_type="application/pdf")


@app.post("/upload_template")
async def upload_template(file: UploadFile = File(...)):
    if not file.filename.endswith(".docx"):
        raise HTTPException(400, "Solo se permiten archivos .docx")

    plantilla_id = str(uuid.uuid4())
    path = os.path.join(TEMPLATES_DIR, file.filename)

    with open(path, "wb") as f:
        f.write(await file.read())

    variables = extraer_variables(path)

    with open(DB_PATH, "r+") as f:
        data = json.load(f)
        data["plantillas"].append({
            "id": plantilla_id,
            "nombre": file.filename,
            "variables": variables
        })
        f.seek(0)
        f.truncate()
        json.dump(data, f, indent=4)

    return {
        "id": plantilla_id,
        "nombre_archivo": file.filename,
        "variables_detectadas": variables
    }


@app.get("/templates")
def list_templates():
    with open(DB_PATH, "r") as f:
        return json.load(f)["plantillas"]


# -------------------------------------------------
# ENDPOINTS DE VARIABLES (con x-api-key)
# -------------------------------------------------
@app.get("/variables")
def get_variables(x_api_key: str = Header(None)):
    if x_api_key != API_ADMIN_KEY:
        raise HTTPException(status_code=401, detail="No autorizado")
    return VARIABLES


@app.put("/variables")
def update_variables(payload: VariablesUpdate, x_api_key: str = Header(None)):
    global VARIABLES

    if x_api_key != API_ADMIN_KEY:
        raise HTTPException(status_code=401, detail="No autorizado")

    data_dict = payload.model_dump(exclude_unset=True)

    # asegurar estructura viva
    if "planes" not in VARIABLES or not isinstance(VARIABLES.get("planes"), dict):
        VARIABLES["planes"] = DEFAULT_VARIABLES["planes"].copy()
    if "defaults" not in VARIABLES or not isinstance(VARIABLES.get("defaults"), dict):
        VARIABLES["defaults"] = DEFAULT_VARIABLES["defaults"].copy()

    # -------- defaults --------
    defaults_in = data_dict.get("defaults")
    if isinstance(defaults_in, dict):
        if "residuales_default" in defaults_in and defaults_in["residuales_default"] is not None:
            defaults_in["residuales_default"] = [
                {"plazo": r["plazo"], "residual": r["residual"]}
                for r in defaults_in["residuales_default"]
            ]

        if "seguro_por_monto" in defaults_in and defaults_in["seguro_por_monto"] is not None:
            defaults_in["seguro_por_monto"] = [
                {"max_valor_con_iva": s["max_valor_con_iva"], "porcentaje": s["porcentaje"]}
                for s in defaults_in["seguro_por_monto"]
            ]

        for k, v in defaults_in.items():
            VARIABLES["defaults"][k] = v

    # -------- planes --------
    planes_in = data_dict.get("planes")
    if isinstance(planes_in, dict):
        for plan_name, plan_updates in planes_in.items():
            pkey = (str(plan_name) or "").strip().upper()
            if not pkey:
                continue

            if pkey not in VARIABLES["planes"] or not isinstance(VARIABLES["planes"].get(pkey), dict):
                VARIABLES["planes"][pkey] = {}

            for k, v in (plan_updates or {}).items():
                VARIABLES["planes"][pkey][k] = v

    # persistir db.json (variables completo)
    with open(DB_PATH, "r+") as f:
        data = json.load(f)
        if "variables" not in data or not isinstance(data["variables"], dict):
            data["variables"] = {}
        data["variables"] = VARIABLES
        f.seek(0)
        f.truncate()
        json.dump(data, f, indent=4)

    refresh_cotizar_example()
    return {"status": "ok", "variables": VARIABLES}

# -------------------------------------------------
# /sync/historico  (Power Automate -> Postgres)
# -------------------------------------------------

_num_re = re.compile(r"[^0-9\.\-]")

_xcode_re = re.compile(r"_x([0-9A-Fa-f]{4})_", re.IGNORECASE)

def _decode_xcodes(s: str) -> str:
    # Convierte _x002e_ -> '.'  y _x0020_ -> ' '
    return _xcode_re.sub(lambda m: chr(int(m.group(1), 16)), s)

def _norm_key(k: str) -> str:
    """
    Normaliza keys provenientes de Excel / Power Automate:
    - Decodifica _xNNNN_ (ej. _x002e_)
    - Quita acentos
    - Unifica separadores
    """
    if k is None:
        return ""
    s = str(k).strip()

    # 🔹 PASO CLAVE: decodificar formato Excel
    s = _decode_xcodes(s)

    # normalización estándar
    s = unicodedata.normalize("NFKD", s)
    s = "".join(ch for ch in s if not unicodedata.combining(ch))
    s = s.replace(".", " ").replace("_", " ")
    s = re.sub(r"\s+", " ", s).strip().upper()
    return s

def _get_any(row: Dict[str, Any], *names: str) -> Any:
    """
    Busca un campo por varios nombres posibles, usando normalización.
    """
    if not row:
        return None
    norm_map = {_norm_key(k): v for k, v in row.items()}
    for n in names:
        v = norm_map.get(_norm_key(n))
        if v is not None and str(v).strip() != "":
            return v
    return None

def _to_number(v: Any):
    if v is None:
        return None
    s = str(v).strip()
    if s == "":
        return None
    s = _num_re.sub("", s)  # quita $, comas, espacios, etc.
    if s in ("", ".", "-", "-.", ".-"):
        return None
    try:
        return float(s)
    except Exception:
        return None

def _to_int(v: Any):
    n = _to_number(v)
    if n is None:
        return None
    try:
        return int(round(n))
    except Exception:
        return None

def _excel_serial_to_date(serial: float):
    # Excel: día 1 = 1899-12-31 (con bug 1900), estándar práctico: base 1899-12-30
    base = datetime(1899, 12, 30)
    return (base + timedelta(days=serial)).date()

def _to_date_str(v: Any):
    if v is None:
        return None
    s = str(v).strip()
    if s == "":
        return None

    # Si viene número de Excel (ej. 43467)
    try:
        n = float(s)
        if n > 20000:  # umbral razonable para fechas modernas
            return _excel_serial_to_date(n).isoformat()
    except Exception:
        pass

    # Si viene ISO o texto, lo devolvemos tal cual
    return s

class SyncTables(BaseModel):
    cartera: List[Dict[str, Any]] = []
    activos: List[Dict[str, Any]] = []

class SyncPayload(BaseModel):
    tables: SyncTables

def _get_cartera_conn():
    if not CARTERA_DATABASE_URL:
        raise HTTPException(status_code=500, detail="Missing CARTERA_DATABASE_URL")
    return psycopg2.connect(CARTERA_DATABASE_URL)

def _norm_str(x) -> Optional[str]:
    if x is None:
        return None
    s = str(x).strip()
    return s or None


def _norm_contrato(x) -> Optional[str]:
    """
    Normaliza el contrato para dedupe:
    - strip
    - deja tal cual el formato, sin inventar guiones.
    """
    s = _norm_str(x)
    if not s:
        return None
    return s


@app.post("/sync/historico")
def sync_historico(payload: SyncPayload, x_api_key: Optional[str] = Header(None)):
    # Seguridad: clave independiente a API_ADMIN_KEY
    if not CARTERA_SYNC_API_KEY:
        raise HTTPException(status_code=500, detail="Missing CARTERA_SYNC_API_KEY")
    if x_api_key != CARTERA_SYNC_API_KEY:
        raise HTTPException(status_code=401, detail="Unauthorized")

    cartera_rows = payload.tables.cartera or []
    activos_rows = payload.tables.activos or []

    # ----------------------------
    # 1) Armar tuples (cartera)
    # ----------------------------
    cartera_map: Dict[str, Tuple] = {}
    cartera_skipped = 0

    for r in cartera_rows:
        contrato = _get_any(r, "NO. CONTRATO", "NO CONTRATO", "NO_CONTRATO", "NO_x002e_ CONTRATO")
        contrato = _norm_contrato(contrato)
        if not contrato:
            cartera_skipped += 1
            continue

        # Normalizados / numéricos
        flujos_inicio_mes = _to_number(_get_any(r, "FLUJOS FUTUROS INICIO MES"))

        # FONDEADOR: solo raw (no columna)
        # Si el Excel manda FONDEADOR, queda guardado en raw automáticamente porque raw=Json(r)

        cartera_map[contrato] = (
            contrato,
            (_get_any(r, "CLIENTE") or None),
            (_get_any(r, "TIPO DE ACTIVO") or None),
            _to_int(_get_any(r, "PLAZO DEL ARRENDAMIENTO")),
            _to_date_str(_get_any(r, "FECHA DE INICIO")),
            _to_date_str(_get_any(r, "FECHA DE VENCIMIENTO")),
            _to_number(_get_any(r, "TASA DE INTERES", "TASA DE INTERÉS")),
            _to_number(_get_any(r, "SALDO INSOLUTO INICIO MES")),
            _to_number(_get_any(r, "PAGOS HISTORICOS C/IVA", "PAGOS HISTÓRICOS C/IVA")),
            flujos_inicio_mes,  # <-- columna normalizada
            Json(r),
        )

    cartera_values = list(cartera_map.values())
    cartera_deduped = len(cartera_rows) - cartera_skipped - len(cartera_values)

    # ----------------------------
    # 2) Armar tuples (activos)
    # ----------------------------
    activos_map: Dict[str, Tuple] = {}
    activos_skipped = 0

    for r in activos_rows:
        ident = _get_any(r, "IDENTIFICADOR", "identificador")
        contrato = _get_any(r, "NO. CONTRATO", "NO CONTRATO", "NO_CONTRATO", "NO_x002e_ CONTRATO")

        ident = _norm_str(ident)
        contrato = _norm_contrato(contrato)

        if not ident or not contrato:
            activos_skipped += 1
            continue

        activos_map[ident] = (
            ident,
            contrato,
            (_get_any(r, "CONTRATO INTERNO") or None),
            (_get_any(r, "CLIENTE") or None),
            (_get_any(r, "TIPO DE ACTIVO") or None),
            (_get_any(r, "DESCRIPCION", "DESCRIPCIÓN") or None),
            (_get_any(r, "NÚMERO DE SERIE", "NUMERO DE SERIE") or None),
            (_get_any(r, "NÚMERO DE MOTOR", "NUMERO DE MOTOR") or None),
            (_to_date_str(_get_any(r, "FECHA DE INICIO"))),
            (_to_date_str(_get_any(r, "FECHA DE VENCIMIENTO"))),
            (_get_any(r, "ASEGURADORA") or None),
            (_get_any(r, "POLIZA", "PÓLIZA") or None),
            (_to_date_str(_get_any(r, "INICIO VIGENCIA POLIZA", "INICIO VIGENCIA PÓLIZA"))),
            (_to_date_str(_get_any(r, "FIN VIGENCIA POLIZA", "FIN VIGENCIA PÓLIZA"))),
            Json(r),
        )

    activos_values = list(activos_map.values())
    activos_deduped = len(activos_rows) - activos_skipped - len(activos_values)

    conn = _get_cartera_conn()
    try:
        with conn:
            with conn.cursor() as cur:
                if cartera_values:
                    execute_values(cur, """
                        insert into public.cartera_historica
                          (no_contrato, cliente, tipo_de_activo, plazo_del_arrendamiento,
                           fecha_de_inicio, fecha_de_vencimiento,
                           tasa_de_interes, saldo_insoluto_inicio_mes, pagos_historicos_c_iva,
                           flujos_futuros_inicio_mes,
                           raw)
                        values %s
                        on conflict (no_contrato) do update set
                          cliente = excluded.cliente,
                          tipo_de_activo = excluded.tipo_de_activo,
                          plazo_del_arrendamiento = excluded.plazo_del_arrendamiento,
                          fecha_de_inicio = excluded.fecha_de_inicio,
                          fecha_de_vencimiento = excluded.fecha_de_vencimiento,
                          tasa_de_interes = excluded.tasa_de_interes,
                          saldo_insoluto_inicio_mes = excluded.saldo_insoluto_inicio_mes,
                          pagos_historicos_c_iva = excluded.pagos_historicos_c_iva,
                          flujos_futuros_inicio_mes = excluded.flujos_futuros_inicio_mes,
                          raw = excluded.raw,
                          updated_at = now();
                    """, cartera_values)

                if activos_values:
                    execute_values(cur, """
                        insert into public.activos_historico
                          (identificador, no_contrato, contrato_interno, cliente, tipo_de_activo,
                           descripcion, numero_de_serie, numero_de_motor,
                           fecha_de_inicio, fecha_de_vencimiento,
                           aseguradora, poliza, inicio_vigencia_poliza, fin_vigencia_poliza,
                           raw)
                        values %s
                        on conflict (identificador) do update set
                          no_contrato = excluded.no_contrato,
                          contrato_interno = excluded.contrato_interno,
                          cliente = excluded.cliente,
                          tipo_de_activo = excluded.tipo_de_activo,
                          descripcion = excluded.descripcion,
                          numero_de_serie = excluded.numero_de_serie,
                          numero_de_motor = excluded.numero_de_motor,
                          fecha_de_inicio = excluded.fecha_de_inicio,
                          fecha_de_vencimiento = excluded.fecha_de_vencimiento,
                          aseguradora = excluded.aseguradora,
                          poliza = excluded.poliza,
                          inicio_vigencia_poliza = excluded.inicio_vigencia_poliza,
                          fin_vigencia_poliza = excluded.fin_vigencia_poliza,
                          raw = excluded.raw,
                          updated_at = now();
                    """, activos_values)

        return {
            "cartera_received": len(cartera_rows),
            "cartera_skipped": cartera_skipped,
            "cartera_deduped": cartera_deduped,
            "cartera_upserted": len(cartera_values),

            "activos_received": len(activos_rows),
            "activos_skipped": activos_skipped,
            "activos_deduped": activos_deduped,
            "activos_upserted": len(activos_values),
        }
    finally:
        conn.close()

# ==============================
# Cartera Query (READ ONLY)
# ==============================
CARTERA_READ_API_KEY = os.getenv("CARTERA_READ_API_KEY")

def _require_read_key(x_api_key: Optional[str]):
    if not CARTERA_READ_API_KEY:
        raise HTTPException(status_code=500, detail="Missing CARTERA_READ_API_KEY")
    if x_api_key != CARTERA_READ_API_KEY:
        raise HTTPException(status_code=401, detail="Unauthorized")

def _clamp_limit(n: Optional[int], default: int = 200, max_n: int = 2000) -> int:
    if n is None:
        return default
    try:
        n = int(n)
    except Exception:
        return default
    return max(1, min(max_n, n))

def _rows_to_dicts(cur):
    cols = [d[0] for d in cur.description]
    return [dict(zip(cols, r)) for r in cur.fetchall()]

_MONTHS_ES = {
    "enero": 1, "febrero": 2, "marzo": 3, "abril": 4, "mayo": 5, "junio": 6,
    "julio": 7, "agosto": 8, "septiembre": 9, "setiembre": 9,
    "octubre": 10, "noviembre": 11, "diciembre": 12,
}

def _month_range_from_q(q: str) -> tuple[date, date]:
    """
    Acepta:
      - 'febrero 2026'
      - 'Feb 2026'
      - '2026-02'
      - '02/2026' o '2/2026'
    Regresa (inicio_mes, inicio_mes_siguiente)
    """
    if not q or not q.strip():
        raise HTTPException(status_code=400, detail="Este scope requiere q con mes y año. Ej: 'febrero 2026' o '2026-02'.")

    s = q.strip().lower()

    # 2026-02
    m = re.match(r"^\s*(\d{4})-(\d{1,2})\s*$", s)
    if m:
        y = int(m.group(1)); mm = int(m.group(2))

    # 02/2026
    elif re.match(r"^\s*\d{1,2}/\d{4}\s*$", s):
        mm, y = s.split("/")
        y = int(y); mm = int(mm)

    # 'febrero 2026' o 'feb 2026'
    else:
        parts = re.split(r"\s+", s)
        if len(parts) < 2:
            raise HTTPException(status_code=400, detail="Formato inválido. Usa 'febrero 2026' o '2026-02'.")
        mes_txt = parts[0][:9]  # tolera abreviaciones
        y = int(parts[1])

        # match por prefijo (feb, febr, febrero)
        mm = None
        for k, v in _MONTHS_ES.items():
            if k.startswith(mes_txt) or mes_txt.startswith(k[:3]):
                if mes_txt.startswith(k[:3]):
                    mm = v
                    break
        if mm is None:
            raise HTTPException(status_code=400, detail="Mes inválido. Ej: 'febrero 2026'.")

    if mm < 1 or mm > 12:
        raise HTTPException(status_code=400, detail="Mes inválido (1-12).")

    start = date(y, mm, 1)
    # siguiente mes
    if mm == 12:
        end = date(y + 1, 1, 1)
    else:
        end = date(y, mm + 1, 1)
    return start, end

class QueryInclude(BaseModel):
    contratos: bool = False
    activos: bool = False

class CarteraQuery(BaseModel):
    scope: str = "cliente"  # cliente | contrato | cartera | activo | fondeador
    q: Optional[str] = None
    no_contrato: Optional[str] = None
    metrics: List[str] = []
    include: QueryInclude = QueryInclude()
    limit: Optional[int] = 200

@app.post("/cartera/query")
def cartera_query(payload: CarteraQuery, x_api_key: Optional[str] = Header(None)):
    _require_read_key(x_api_key)

    scope = (payload.scope or "cliente").strip().lower()
    q = (payload.q or "").strip()
    no_contrato = (payload.no_contrato or "").strip()
    limit = _clamp_limit(payload.limit)

    # -----------------------------
    # Filtros base (cartera)
    # -----------------------------
    where_cartera = "1=1"
    params_cartera: list = []

    # -----------------------------
    # Filtros activos (solo se usan si scope=activo)
    # -----------------------------
    where_activos: Optional[str] = None
    params_activos: Optional[list] = None

    if scope == "cliente":
        if not q:
            raise HTTPException(status_code=400, detail="scope=cliente requires q (cliente)")
        where_cartera = "c.cliente ILIKE %s"
        params_cartera = [f"%{q}%"]

    elif scope == "contrato":
        if not no_contrato:
            no_contrato = q
        if not no_contrato:
            raise HTTPException(status_code=400, detail="scope=contrato requires no_contrato (or q)")

        no_contrato_clean = no_contrato.strip()

        # Si son 4 dígitos, interpretarlo como el 3er bloque del contrato: XXXXX-XXXXX-0166-XXX
        if len(no_contrato_clean) == 4 and no_contrato_clean.isdigit():
            where_cartera = "c.no_contrato ILIKE %s"
            params_cartera = [f"%-{no_contrato_clean}-%"]
        else:
            where_cartera = "c.no_contrato = %s"
            params_cartera = [no_contrato_clean]

        # 🔎 PRECHECK: validar que el contrato exista
        with _get_cartera_conn() as conn_check:
            with conn_check.cursor() as cur_check:
                cur_check.execute(
                    f"SELECT 1 FROM public.cartera_historica c WHERE {where_cartera} LIMIT 1;",
                    params_cartera
                )
                if cur_check.fetchone() is None:
                        return {
                            "ok": False,
                            "error": {
                                "code": "CONTRACT_NOT_FOUND",
                                "detail": "Contrato no encontrado con el identificador proporcionado"
                            },
                            "filters_applied": {"scope": scope, "q": q or None, "no_contrato": no_contrato_clean},
                            "metrics": {},
                            "rows": {"contratos": [], "activos": []},
                        }

    elif scope == "cartera":
        where_cartera = "1=1"
        params_cartera = []
    elif scope == "vencimientos_contrato":
        # q requerido: "febrero 2026" o "2026-02"
        start, end = _month_range_from_q(q)

        # filtro sobre cartera_historica
        where_cartera = "c.fecha_de_vencimiento >= %s AND c.fecha_de_vencimiento < %s"
        params_cartera = [start, end]

    elif scope == "vencimientos_seguro":
        # q requerido: "marzo 2026" o "2026-03"
        start, end = _month_range_from_q(q)

        # aquí NO usamos where_cartera: vamos a listar activos por fin_vigencia_poliza
        where_activos = "a.fin_vigencia_poliza >= %s AND a.fin_vigencia_poliza < %s"
        params_activos = [start, end]

    elif scope == "fondeador":
        # Si viene q => filtra por ese fondeador (match parcial)
        # Si NO viene q => solo “contratos que tienen fondeador”
        if q:
            where_cartera = "NULLIF(BTRIM(c.fondeador), '') IS NOT NULL AND c.fondeador ILIKE %s"
            params_cartera = [f"%{q}%"]
        else:
            where_cartera = "NULLIF(BTRIM(c.fondeador), '') IS NOT NULL"
            params_cartera = []

    # ---------------------------------------------------------
    # --- FONDEADOR SCOPE ---
    #
    # Si NO viene q -> “contratos que tienen fondeador”
    # Si viene q -> “contratos fondeados por <q>”
    # ---------------------------------------------------------
    elif scope == "fondeador":
        # condición SQL: fondeador no nulo y no vacío
        fondeador_has_value_sql = "NULLIF(BTRIM(c.fondeador), '') IS NOT NULL"

        if q:
            where_cartera = f"({fondeador_has_value_sql} AND c.fondeador ILIKE %s)"
            params_cartera = [f"%{q}%"]
        else:
            where_cartera = f"({fondeador_has_value_sql})"
            params_cartera = []

    elif scope == "activo":
        if not q:
            raise HTTPException(status_code=400, detail="scope=activo requires q")

        q_clean = q.strip()

        # 1) Si parece contrato completo (tiene 3+ guiones), match exacto por no_contrato
        if q_clean.count("-") >= 3:
            where_activos = "a.no_contrato = %s"
            params_activos = [q_clean]

        # 2) Si son 4 dígitos, match por 3er bloque (ignora ceros a la izquierda)
        elif len(q_clean) == 4 and q_clean.isdigit():
            where_activos = "ltrim(split_part(a.no_contrato, '-', 3), '0') = ltrim(%s, '0')"
            params_activos = [q_clean]

        # 3) Búsqueda amplia (texto libre)
        else:
            like = f"%{q_clean}%"
            where_activos = """
            (
                a.no_contrato ILIKE %s OR
                coalesce(a.contrato_interno,'') ILIKE %s OR
                a.cliente ILIKE %s OR
                a.tipo_de_activo ILIKE %s OR
                a.descripcion ILIKE %s OR
                a.numero_de_serie ILIKE %s OR
                a.numero_de_motor ILIKE %s OR
                a.aseguradora ILIKE %s OR
                a.poliza ILIKE %s
            )
            """
            params_activos = [like, like, like, like, like, like, like, like, like]

    else:
        raise HTTPException(status_code=400, detail="Invalid scope")

    # Helpers SQL para extraer numericos desde raw (tolerante a comas/$)
    raw_num = lambda key: f"""nullif(regexp_replace(coalesce(c.raw->>'{key}',''), '[^0-9\\.\\-]', '', 'g'),'')::numeric"""
    raw_int = lambda key: f"""nullif(regexp_replace(coalesce(c.raw->>'{key}',''), '[^0-9\\-]', '', 'g'),'')::int"""

    metrics_out = {}
    conn = _get_cartera_conn()
    try:
        with conn, conn.cursor() as cur:

            # ---------- MÉTRICAS ----------
            for m in (payload.metrics or []):
                m = (m or "").strip().lower()

                if m == "conteo_contratos":
                    cur.execute(
                        f"select count(*) as conteo_contratos from public.cartera_historica c where {where_cartera};",
                        params_cartera
                    )
                    metrics_out["conteo_contratos"] = int(cur.fetchone()[0] or 0)

                elif m == "vigentes_terminados":
                    cur.execute(f"""
                        select
                          count(*) filter (where coalesce(c.saldo_insoluto_inicio_mes,0) > 0) as vigentes,
                          count(*) filter (where coalesce(c.saldo_insoluto_inicio_mes,0) = 0) as terminados
                        from public.cartera_historica c
                        where {where_cartera};
                    """, params_cartera)
                    r = cur.fetchone()
                    metrics_out["vigentes_terminados"] = {"vigentes": int(r[0] or 0), "terminados": int(r[1] or 0)}

                elif m == "suma_rentas_vigentes":
                    cur.execute(f"""
                        select coalesce(sum({raw_num('FLUJO MENSUAL S/IVA')}),0) as suma_rentas_vigentes
                        from public.cartera_historica c
                        where {where_cartera}
                          and coalesce(c.saldo_insoluto_inicio_mes,0) > 0;
                    """, params_cartera)
                    metrics_out["suma_rentas_vigentes"] = float(cur.fetchone()[0] or 0)

                elif m == "suma_cartera":
                    cur.execute(f"""
                        select coalesce(sum(coalesce(c.flujos_futuros_inicio_mes, 0)), 0) as suma_cartera
                        from public.cartera_historica c
                        where {where_cartera};
                    """, params_cartera)
                    metrics_out["suma_cartera"] = float(cur.fetchone()[0] or 0)

                elif m == "suma_pagos":
                    cur.execute(f"""
                        select coalesce(sum(coalesce(c.pagos_historicos_c_iva,0)),0) as suma_pagos
                        from public.cartera_historica c
                        where {where_cartera};
                    """, params_cartera)
                    metrics_out["suma_pagos"] = float(cur.fetchone()[0] or 0)

                elif m == "saldo_insoluto":
                    cur.execute(f"""
                        select coalesce(sum(coalesce(c.saldo_insoluto_inicio_mes,0)),0) as saldo_insoluto
                        from public.cartera_historica c
                        where {where_cartera};
                    """, params_cartera)
                    metrics_out["saldo_insoluto"] = float(cur.fetchone()[0] or 0)

                elif m == "rentas_por_devengar":
                    if scope == "contrato":
                        cur.execute(f"""
                            select c.no_contrato,
                                   {raw_int('NO. DE RENTAS POR DEVENGAR')} as rentas_por_devengar
                            from public.cartera_historica c
                            where {where_cartera}
                            limit 1;
                        """, params_cartera)
                        row = cur.fetchone()
                        metrics_out["rentas_por_devengar"] = (int(row[1]) if row and row[1] is not None else None)
                    else:
                        cur.execute(f"""
                            select coalesce(sum(coalesce({raw_int('NO. DE RENTAS POR DEVENGAR')},0)),0) as rentas_por_devengar
                            from public.cartera_historica c
                            where {where_cartera};
                        """, params_cartera)
                        metrics_out["rentas_por_devengar"] = int(cur.fetchone()[0] or 0)

                elif m == "monto_financiado":
                    cur.execute(f"""
                        select coalesce(sum({raw_num('MONTO A FINANCIAR S/IVA')}),0) as monto_financiado
                        from public.cartera_historica c
                        where {where_cartera};
                    """, params_cartera)
                    metrics_out["monto_financiado"] = float(cur.fetchone()[0] or 0)

                elif m == "valor_residual":
                    cur.execute(f"""
                        select coalesce(sum({raw_num('VALOR RESIDUAL S/IVA')}),0) as valor_residual
                        from public.cartera_historica c
                        where {where_cartera};
                    """, params_cartera)
                    metrics_out["valor_residual"] = float(cur.fetchone()[0] or 0)

                elif m == "condiciones":
                    if scope != "contrato":
                        metrics_out["condiciones"] = "scope=contrato required"
                    else:
                        cur.execute(f"""
                            select
                              c.no_contrato,
                              c.plazo_del_arrendamiento,
                              c.tasa_de_interes,
                              {raw_num('APORTACION EXTRAORDINARIA S/IVA')} as aportacion_extraordinaria_s_iva,
                              {raw_num('COMISION POR APERTURA S/IVA')} as comision_por_apertura_s_iva,
                              {raw_num('DEPOSITO EN GARANTIA S/IVA')} as deposito_en_garantia_s_iva,
                              {raw_num('MONTO A FINANCIAR S/IVA')} as monto_a_financiar_s_iva,
                              {raw_num('VALOR FACTURA ACTIVO S/IVA')} as valor_factura_activo_s_iva,
                              {raw_num('VALOR RESIDUAL S/IVA')} as valor_residual_s_iva
                            from public.cartera_historica c
                            where {where_cartera}
                            limit 1;
                        """, params_cartera)
                        r = cur.fetchone()
                        if not r:
                            metrics_out["condiciones"] = None
                        else:
                            metrics_out["condiciones"] = {
                                "no_contrato": r[0],
                                "plazo": r[1],
                                "tasa_interes": float(r[2]) if r[2] is not None else None,
                                "aportacion_extraordinaria_s_iva": float(r[3]) if r[3] is not None else None,
                                "comision_por_apertura_s_iva": float(r[4]) if r[4] is not None else None,
                                "deposito_en_garantia_s_iva": float(r[5]) if r[5] is not None else None,
                                "monto_a_financiar_s_iva": float(r[6]) if r[6] is not None else None,
                                "valor_factura_activo_s_iva": float(r[7]) if r[7] is not None else None,
                                "valor_residual_s_iva": float(r[8]) if r[8] is not None else None,
                            }

                elif m == "fechas":
                    if scope == "contrato":
                        cur.execute(f"""
                            select c.no_contrato, c.fecha_de_inicio, c.fecha_de_vencimiento
                            from public.cartera_historica c
                            where {where_cartera}
                            limit 1;
                        """, params_cartera)
                        r = cur.fetchone()
                        metrics_out["fechas"] = {
                            "no_contrato": r[0] if r else None,
                            "fecha_inicio": (r[1].isoformat() if r and r[1] else None),
                            "fecha_vencimiento": (r[2].isoformat() if r and r[2] else None),
                        }
                    else:
                        cur.execute(f"""
                            select min(c.fecha_de_inicio) as min_inicio,
                                   max(c.fecha_de_inicio) as max_inicio,
                                   min(c.fecha_de_vencimiento) as min_venc,
                                   max(c.fecha_de_vencimiento) as max_venc
                            from public.cartera_historica c
                            where {where_cartera};
                        """, params_cartera)
                        r = cur.fetchone()
                        metrics_out["fechas"] = {
                            "min_inicio": r[0].isoformat() if r and r[0] else None,
                            "max_inicio": r[1].isoformat() if r and r[1] else None,
                            "min_vencimiento": r[2].isoformat() if r and r[2] else None,
                            "max_vencimiento": r[3].isoformat() if r and r[3] else None,
                        }
                # --- FONDEADOR: “Cuántos contratos tienen fondeador” ---
                elif m == "conteo_contratos_con_fondeador":
                    cur.execute("""
                        select count(distinct c.no_contrato) as conteo_con_fondeador
                        from public.cartera_historica c
                        where NULLIF(BTRIM(c.fondeador), '') IS NOT NULL;
                    """)
                    metrics_out["conteo_contratos_con_fondeador"] = int(cur.fetchone()[0] or 0)

                # --- FONDEADOR: “Qué fondeadores existen” ---
                elif m == "lista_fondeadores":
                    cur.execute("""
                        select distinct BTRIM(c.fondeador) as fondeador
                        from public.cartera_historica c
                        where NULLIF(BTRIM(c.fondeador), '') IS NOT NULL
                        order by fondeador asc
                        limit 200;
                    """)
                    metrics_out["lista_fondeadores"] = [r[0] for r in cur.fetchall() if r and r[0]]

                elif m == "desde_cuando_cliente":
                    cur.execute(f"""
                        select min(c.fecha_de_inicio) as desde_cuando
                        from public.cartera_historica c
                        where {where_cartera};
                    """, params_cartera)
                    r = cur.fetchone()
                    metrics_out["desde_cuando_cliente"] = (r[0].isoformat() if r and r[0] else None)

                else:
                    metrics_out[m] = "unknown_metric"

            # ---------- LISTADOS ----------
            out_rows = {"contratos": [], "activos": []}

            if payload.include and payload.include.contratos:
                cur.execute(f"""
                    select
                      c.no_contrato, c.cliente, c.tipo_de_activo, c.fondeador,
                      c.plazo_del_arrendamiento, c.fecha_de_inicio, c.fecha_de_vencimiento,
                      c.saldo_insoluto_inicio_mes, c.pagos_historicos_c_iva,
                      c.raw->>'FLUJO MENSUAL S/IVA' as flujo_mensual_s_iva,
                      c.raw->>'FLUJOS FUTUROS INICIO MES' as flujos_futuros_inicio_mes
                    from public.cartera_historica c
                    where {where_cartera}
                    order by c.updated_at desc
                    limit %s;
                """, params_cartera + [limit])
                out_rows["contratos"] = _rows_to_dicts(cur)

            # Mostrar activos cuando:
            # - scope sea "activo" (siempre regresa activos)
            # - o include.activos = true
            want_activos = (scope in ("activo", "vencimientos_seguro")) or (payload.include and payload.include.activos)

            if want_activos:
                if scope in ("activo", "vencimientos_seguro"):
                    if where_activos is None or params_activos is None:
                        raise HTTPException(status_code=400, detail="Este scope requiere q para filtrar seguros/activos")
                    cur.execute(f"""
                        select
                            a.identificador, a.no_contrato, a.cliente, a.tipo_de_activo, a.descripcion,
                            a.numero_de_serie, a.numero_de_motor, a.aseguradora, a.poliza,
                            a.inicio_vigencia_poliza, a.fin_vigencia_poliza
                        from public.activos_historico a
                        where {where_activos}
                        order by a.fin_vigencia_poliza asc, a.no_contrato asc, a.identificador asc
                        limit %s;
                    """, params_activos + [limit])
                    out_rows["activos"] = _rows_to_dicts(cur)

                elif scope == "contrato":
                    nc = (no_contrato or q or "").strip()
                    if not nc:
                        raise HTTPException(status_code=400, detail="include.activos with scope=contrato requires no_contrato (or q)")

                    if len(nc) == 4 and nc.isdigit():
                        cur.execute("""
                            select identificador, no_contrato, cliente, tipo_de_activo, descripcion,
                                   numero_de_serie, numero_de_motor, aseguradora, poliza,
                                   inicio_vigencia_poliza, fin_vigencia_poliza
                            from public.activos_historico
                            where ltrim(split_part(no_contrato,'-',3),'0') = ltrim(%s,'0')
                            order by identificador
                            limit %s;
                        """, (nc, limit))
                    else:
                        cur.execute("""
                            select identificador, no_contrato, cliente, tipo_de_activo, descripcion,
                                   numero_de_serie, numero_de_motor, aseguradora, poliza,
                                   inicio_vigencia_poliza, fin_vigencia_poliza
                            from public.activos_historico
                            where no_contrato = %s
                            order by identificador
                            limit %s;
                        """, (nc, limit))
                    out_rows["activos"] = _rows_to_dicts(cur)

                else:
                    # otros scopes: solo si hay q (búsqueda simple)
                    if q:
                        cur.execute("""
                            select identificador, no_contrato, cliente, tipo_de_activo, descripcion,
                                   numero_de_serie, numero_de_motor, aseguradora, poliza,
                                   inicio_vigencia_poliza, fin_vigencia_poliza
                            from public.activos_historico
                            where descripcion ilike %s
                               or numero_de_serie ilike %s
                               or numero_de_motor ilike %s
                               or poliza ilike %s
                            order by a.updated_at desc
                            limit %s;
                        """, (f"%{q}%", f"%{q}%", f"%{q}%", f"%{q}%", limit))
                        out_rows["activos"] = _rows_to_dicts(cur)

        return {
            "filters_applied": {"scope": scope, "q": q or None, "no_contrato": no_contrato or None},
            "metrics": metrics_out,
            "rows": out_rows
        }
    finally:
        conn.close()

# -------------------------------------------------
# HEALTH CHECK
# -------------------------------------------------
@app.get("/")
def health_check():
    return {"status": "ok"}
