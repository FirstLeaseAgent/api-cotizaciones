from docx import Document
import re

def extraer_variables(docx_path):
    """
    Escanea el documento .docx y devuelve una lista de variables entre {{ }}
    """
    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"Error al leer el documento: {e}")
        return []

    texto = ""
    for p in doc.paragraphs:
        texto += p.text + " "
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texto += cell.text + " "

    # ---Busca Variables dentro {{variable}}----
    variables = re.findall(r"\{\{(.*?)\}\}", texto)
    return list(set(v.strip() for v in variables))
